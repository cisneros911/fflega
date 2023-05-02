import openai
import os
import time
from docx import Document
from docx.shared import Inches
import requests
import json

api_key = "YOUR_OPENAI_API_KEY"
openai.api_key = api_key

def obtener_ideas():
    prompt_text = f"Dame un título de cuento para niños"

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt_text,
        max_tokens=500,
        n=1,
        temperature=0.5,
        top_p=1
    )

    ideas = response.choices[0].text.strip()
    return ideas


cuantos_cuentos = input("¿Cuántos cuentos para niños quieres generar? (Por defecto: '5'): ")
if not cuantos_cuentos:
    cuantos_cuentos = 5
else:
    cuantos_cuentos = int(cuantos_cuentos)

titulos_cuentos = []
for _ in range(cuantos_cuentos):
    titulo = obtener_ideas()
    titulos_cuentos.append(titulo)

while True:
    respuesta = input(f"Desea comenzar un libro con el siguiente contenido? {', '.join(titulos_cuentos)} (Y/N, por defecto Y): ")
    if not respuesta or respuesta.lower() == 'y':
        break
    elif respuesta.lower() == 'n':
        print("Operación cancelada por el usuario.")
        exit()

def generar_cuento(titulo, texto_previo=""):
    prompt = f"Escribe un cuento para niños titulado '{titulo}' que termine con 'Fin.'"
    if texto_previo:
        prompt = f"{texto_previo}"
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=500,
        n=1,
        #stop="Fin.",
        temperature=0.5,
        top_p=1
    )
    cuento = response.choices[0].text.strip()
    #if cuento.endswith("Fin."):
    return cuento
    #else:
    #    texto_completo = f"{texto_previo}{cuento}"
    #    return generar_cuento(titulo, texto_completo)

def generar_imagen(prompt):
    response = openai.Image.create(
        prompt=prompt,
        n=1,
        size="1024x1024"
    )
    image_url = response['data'][0]['url']
    return image_url

def agregar_imagen_docx(doc, url_imagen):
    response = requests.get(url_imagen)
    imagen_path = 'imagen_temporal.jpg'
    with open(imagen_path, 'wb') as f:
        f.write(response.content)

    doc.add_picture(imagen_path, width=Inches(4.0))
    os.remove(imagen_path)

def guardar_cuentos_en_docx(titulos, archivo_salida):
    doc = Document()
    total_cuentos = len(titulos)
    tiempo_inicial = time.time()

    for i, titulo in enumerate(titulos):
        cuento = generar_cuento(titulo)
        doc.add_heading(titulo, level=1)
        doc.add_paragraph(cuento)
        
        imagen_prompt = f"Una imagen para ilustrar un cuento de niños titulado '{titulo}'"
        url_imagen = generar_imagen(imagen_prompt)
        agregar_imagen_docx(doc, url_imagen)
        
        doc.add_page_break()

        # Mostrar el progreso
        progreso = (i + 1) / total_cuentos
        tiempo_actual = time.time()
        tiempo_transcurrido = tiempo_actual - tiempo_inicial
        tiempo_restante = tiempo_transcurrido * (1 / progreso - 1)

        print(
            f"Progreso: {progreso * 100:.2f}% - Tiempo restante estimado: {tiempo_restante:.2f} segundos"
        )
        
        # Guardar el archivo docx en cada iteración para agregar cuentos uno tras otro
        doc.save(archivo_salida)

        # Pausa de 30 segundos después de cada 26 cuentos
        if (i + 1) % 26 == 0:
            time.sleep(30)

# Crear un nombre de archivo único basado en la entrada del usuario y la fecha y hora de creación
#nombre_limpio = tema_libro.replace(" ", "-").lower()
#tiempo_actual = time.strftime("%Y-%m-%d_%I:%M%p")
#archivo_salida = f"{nombre_limpio}-{tiempo_actual}.docx"
archivo_salida = f"book.docx"
#archivo_salida = f"{tiempo_actual}.docx"

guardar_cuentos_en_docx(titulos_cuentos, archivo_salida)
print(f"Se ha creado el documento '{archivo_salida}'")
